import json
from collections import OrderedDict
import pandas as pd
from typing import Optional
import os
from pathlib import Path
from .myconfig import config_manager
from app.log import logger

from .content import Chapter
from .feature import Feature
from .tool import get_description, get_structure, get_flow_chart

from docx import Document
from docx.shared import Inches
import concurrent.futures
from functools import partial

class MetaDataProcessor:
    """处理Excel元数据的类"""
    
    def __init__(self):
        """初始化MetaDataProcessor，从配置文件读取参数"""
        config = config_manager.config
        self.file_name = Path("web/public/resource/file/uploads") / config["excel_file_name"]
        self.sheet_name = config["sheet_name"]
        self.df: Optional[pd.DataFrame] = None
        # self.columns = ["B", "C", "D", "E", "F", "H", "L"]
        self.columns = ['功能用户需求', '触发事件', '功能过程', '子过程描述', '数据组', '功能用户', '角色']
        self.output_path = Path("web/public/resource/file/output/data.json")
        
        self.progress = {
            "current_key": "",
            "total": 0,
            "current": 0,
            "percentage": 0
        }

    def _clean_text(self, text: str) -> str:
        """清理文本中的空白字符"""
        return str(text).strip().replace(" ", "").replace("\t", "").replace("\n", "")

    def read_excel(self) -> pd.DataFrame:
        """读取Excel文件并缓存结果"""
        if self.df is None:
            try:
                self.df = pd.read_excel(
                    self.file_name,
                    sheet_name=self.sheet_name,
                    header=0,
                    usecols=self.columns
                ).fillna(method='ffill')
                logger.info(f"成功读取Excel文件: {self.file_name}")
            except FileNotFoundError:
                logger.error(f"文件未找到: {self.file_name}")
                raise
            except Exception as e:
                logger.error(f"读取Excel文件时出错: {str(e)}")
                raise
        return self.df

    def check_info(self, json_path: Path):
        """
        校验生成的JSON文件内容
        
        Args:
            json_path: JSON文件路径
            
        Returns:
            problems: 发现的问题数量列表
        """
        with open(json_path, 'r', encoding='utf-8') as file:
            data = json.load(file)
        problems = []

        for key, value in data.items():
            # 功能
            functions = []
            # 特性组
            features = []
            for item_index, (item_key, item_value) in enumerate(value.items()):
                if item_index == 0:  # 从第二项开始
                    item_value = item_value.replace(',', '，')
                    role = item_value.strip().split("，")
                    if len(role) != 3:
                        problems.append(f"角色信息不全，{key}")
                        logger.error(f"角色信息不全，{key}")
                else:
                    functions.append(item_key)
                    scenario = item_key
                    process = item_value[0]
                    if len(process) != 3:
                        pass
                        # problems += 1
                        problems.append(f"角色信息不全，{key}")
                        logger.error(f"功能过程不是三个，{key}")
        return problems

    def process_data(self):
        try:
            # 重置 DataFrame 缓存，确保读取最新的 Excel 文件
            self.df = None
            
            # 初始化进度
            self.progress = {
                "current_key": "正在解析Excel数据",
                "total": 0,
                "current": 0,
                "percentage": 0
            }
            
            df = self.read_excel()
            result_dict = OrderedDict()

            # 处理每一行数据
            for _, row in df.iterrows():
                func_user_req = row['功能用户需求']
                trigger_event = row['触发事件']
                func_process = row['功能过程']
                sub_processes = row['子过程描述']
                data_group = row['数据组']
                # fun_user = row['功能用户']
                role = self._clean_text(row['角色'])

                # 初始化功能用户需求字典
                result_dict.setdefault(func_user_req, OrderedDict())
                
                # 设置角色
                result_dict[func_user_req].setdefault("角色", role)
                
                # 初始化功能过程
                result_dict[func_user_req].setdefault(func_process, [[], []])
                
                # 添加子过程和数据组
                sub_proc_list, data_group_list = result_dict[func_user_req][func_process]
                if sub_processes not in sub_proc_list:
                    sub_proc_list.append(sub_processes)
                if data_group not in data_group_list:
                    data_group_list.append(data_group)

            # 转换为JSON
            json_result = json.dumps(result_dict, ensure_ascii=False, indent=4)
            
          
            with open(self.output_path, 'w', encoding='utf-8') as file:
                file.write(json_result)
                logger.info(f"数据已保存至: {self.output_path}")
            return json_result

        except Exception as e:
            logger.error(f"处理数据时出错: {str(e)}")
            raise
    
    def process_feature(self, item_key, item_value, role):
        """处理单个特性的辅助函数"""
        scenario = item_key
        process = item_value[0]
        flow_chart = get_flow_chart(role, process)
        input_data = item_value[1][0]
        output = item_value[1][-1]
        return Feature(scenario, flow_chart, process, input_data, output, role)

    def get_data_result(self) -> list:
        """
        并行处理数据封装
        
        Returns:
            list: 包含 Chapter 对象的列表
        """
        try:
            with open(self.output_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
            result = []

            total_items = len(data.items())
            self.progress["total"] = total_items
            
            # 创建线程池执行器
            with concurrent.futures.ThreadPoolExecutor(max_workers=8) as executor:
                for idx, (key, value) in enumerate(data.items(), 1):
                    self.progress.update({
                        "current_key": key,
                        "current": idx,
                        "percentage": int((idx / total_items) * 100)
                    })
                    logger.info(f"正在处理: {key} ({idx}/{total_items})")
                    
                    functions = []
                    features = []
                    role = None
                    
                    # 收集需要并行处理的特性任务
                    feature_tasks = []
                    for item_index, (item_key, item_value) in enumerate(value.items()):
                        if item_index == 0:
                            item_value = item_value.replace(',', '，')
                            role = item_value.strip().split("，")
                        else:
                            functions.append(item_key)
                            # 创建特性处理任务
                            feature_tasks.append(
                                executor.submit(
                                    self.process_feature,
                                    item_key,
                                    item_value,
                                    role
                                )
                            )
                    
                    # 收集所有特性处理结果
                    for future in concurrent.futures.as_completed(feature_tasks):
                        try:
                            feature = future.result()
                            features.append(feature)
                        except Exception as e:
                            logger.error(f"处理特性时出错: {str(e)}")
                            raise

                    # 生成章节描述和结构（这部分保持串行）
                    description = get_description(key, functions)
                    structure = get_structure(key, functions)
                    chapter = Chapter(
                        name=key,
                        description=description,
                        functions=functions,
                        structure=structure,
                        feature=features
                    )
                    result.append(chapter)
                
            logger.info("数据封装完成")
            return result
            
        except Exception as e:
            logger.error(f"数据封装过程中出错: {str(e)}")
            raise

    def get_docx(self, result: list, output_file: str = "web/public/resource/file/output/data.docx") -> None:
        try:
            doc = Document()
            
            for index, item in enumerate(result, start=1):
                # 章节标题
                doc.add_heading(f'{index}. {item.name}', level=1)
                
                # 1. 产品概述
                doc.add_heading(f'{index}.1. 产品概述', level=2)
                doc.add_paragraph(item.description)

                # 2. 产品结构（功能摘要)
                doc.add_heading(f'{index}.2. 产品结构（功能摘要)', level=2)
                doc.add_paragraph('产品结构如图：')
                # 修改这里：将 Path 对象转换为字符串
                doc.add_picture(str(item.structure), width=Inches(7))
                
                # 功能列表
                doc.add_paragraph('主要包括如下功能')
                for fun in item.functions:
                    doc.add_paragraph(fun, style='List Bullet')

                # 3. 特性说明
                doc.add_heading(f'{index}.3. 特性说明', level=2)

                for loc, detail in enumerate(item.feature, start=1):
                    # 特性子标题
                    doc.add_heading(f'{index}.3.{loc}. {detail.scenario}', level=3)
                    doc.add_paragraph(f'用户场景： {detail.scenario}', style='List Bullet')
                    
                    # 流程图
                    doc.add_paragraph('流程图', style='List Bullet')
                    # 修改这里：将 Path 对象转换为字符串
                    doc.add_picture(str(detail.flow_chart), width=Inches(7))
                    
                    # 功能过程
                    doc.add_paragraph('功能过程:', style='List Bullet')
                    for step, order in enumerate(detail.process, start=1):
                        # 使用 rstrip() 去除尾部空白字符，避免多余空行
                        doc.add_paragraph(f'{step}. {order}'.rstrip())
                    
                    # 输入输出
                    doc.add_paragraph(f'输入：{detail.input}', style='List Bullet')
                    doc.add_paragraph(f'输出：{detail.output}', style='List Bullet')

            # 保存文档
            doc.save(output_file)
            logger.info(f"Word文档已生成：{output_file}")
            
        except Exception as e:
            logger.error(f"生成Word文档时出错: {str(e)}")
            raise
    
    
    async def get_progress(self):
        """获取当前处理进度"""
        return self.progress
resolver = MetaDataProcessor()